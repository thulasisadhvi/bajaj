import pdfplumber
from docx import Document
import email
from email import policy
from email.parser import BytesParser
# import requests # <-- REMOVE THIS LINE
import httpx # <-- ADD THIS LINE
import io
import mimetypes
from bs4 import BeautifulSoup
import html2text
import os # Add os for test_reader cleanup if you run __main__ block

def _extract_text_from_pdf(file_content: bytes) -> str:
    """Extracts text from PDF bytes, including basic table extraction."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Basic table extraction and formatting
                tables = page.extract_tables()
                for table in tables:
                    table_str = "\nTable:\n"
                    for row in table:
                        # Join cells with a tab or pipe for readability
                        table_str += "\t".join([cell if cell is not None else "" for cell in row]) + "\n"
                    text_parts.append(table_str)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        # Consider raising or logging the full traceback in development
        raise # <-- Temporarily raise to see the full error if it happens here
        # return "" # <-- Keep this in production

def _extract_text_from_docx(file_content: bytes) -> str:
    """Extracts text from DOCX bytes, including basic table extraction."""
    full_text = []
    try:
        document = Document(io.BytesIO(file_content))
        for element in document.element.body:
            if element.tag.endswith('p'): # Paragraph
                full_text.append(element.text)
            elif element.tag.endswith('tbl'): # Table
                # Create a temp doc from table XML
                # Note: docx.Document(element.xml) is not standard.
                # You usually access tables via document.tables
                # A more robust table extraction would involve iterating document.tables directly.
                # For now, keeping original logic, but be aware.
                table_element = Document(io.BytesIO(element.xml.encode('utf-8'))) # Ensure bytes for io.BytesIO
                for table in table_element.tables:
                    table_str = "\nTable:\n"
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = ""
                            for paragraph in cell.paragraphs:
                                cell_text += paragraph.text
                            row_cells.append(cell_text)
                        table_str += "\t".join(row_cells) + "\n"
                    full_text.append(table_str)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        raise # <-- Temporarily raise
        # return ""

def _extract_text_from_email_bytes(file_content: bytes) -> str:
    """Extracts text content from email bytes."""
    msg = None
    try:
        msg = BytesParser(policy=policy.default).parsebytes(file_content)

        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                cdispo = str(part.get('Content-Disposition'))
                if ctype == 'text/plain' and 'attachment' not in cdispo:
                    return part.get_payload(decode=True).decode(errors='ignore')
                elif ctype == 'text/html' and 'attachment' not in cdispo:
                    h = html2text.HTML2Text()
                    h.ignore_links = False
                    h.ignore_images = True
                    return h.handle(part.get_payload(decode=True).decode(errors='ignore'))
        else:
            if msg.get_content_type() == 'text/html':
                h = html2text.HTML2Text()
                h.ignore_links = False
                h.ignore_images = True
                return h.handle(msg.get_payload(decode=True).decode(errors='ignore'))
            return msg.get_payload(decode=True).decode(errors='ignore')
    except Exception as e:
        print(f"Error extracting text from email: {e}")
        raise # <-- Temporarily raise
        # return ""

def _extract_text_from_html(file_content: bytes) -> str:
    """Extracts text from HTML bytes, converting to Markdown-like text."""
    try:
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True
        return h.handle(file_content.decode(errors='ignore'))
    except Exception as e:
        print(f"Error extracting text from HTML: {e}")
        raise # <-- Temporarily raise
        # return ""

async def read_document_from_url(document_url: str) -> str:
    """
    Fetches document content from a URL and extracts text based on content type.
    """
    try:
        # Use httpx.AsyncClient for asynchronous requests
        # It's good practice to create and reuse a single AsyncClient instance for a FastAPI app
        # (e.g., in app startup/shutdown events), but for a single function call, this is fine.
        async with httpx.AsyncClient(timeout=30.0) as client: # Added timeout
            response = await client.get(document_url, follow_redirects=True) # Use await client.get
            response.raise_for_status() # Raise an exception for 4xx/5xx responses
            content_type = response.headers.get("Content-Type", "").split(';')[0].strip().lower()
            file_content = response.content # httpx response.content is bytes

        print(f"Detected Content-Type: {content_type}")

        if "application/pdf" in content_type:
            return _extract_text_from_pdf(file_content)
        elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type:
            return _extract_text_from_docx(file_content)
        elif "message/rfc822" in content_type or document_url.lower().endswith(".eml"):
            return _extract_text_from_email_bytes(file_content)
        elif "text/html" in content_type:
            return _extract_text_from_html(file_content)
        elif "text/plain" in content_type:
            return file_content.decode(errors='ignore')
        else:
            print(f"Warning: Unknown content type '{content_type}'. Attempting to decode as plain text.")
            return file_content.decode(errors='ignore')

    except httpx.RequestError as e: # Catch httpx-specific request errors (e.g., network issues, timeouts)
        print(f"Network or request error fetching {document_url} with httpx: {e}")
        raise # <-- TEMPORARILY RAISE FOR DEBUGGING
    except httpx.HTTPStatusError as e: # Catch HTTP 4xx/5xx errors
        print(f"HTTP error fetching {document_url}: {e.response.status_code} - {e.response.text}")
        raise # <-- TEMPORARILY RAISE FOR DEBUGGING
    except Exception as e: # Catch any other unexpected errors
        print(f"General error reading document from {document_url}: {e}")
        raise # <-- TEMPORARILY RAISE FOR DEBUGGING
    # In a production environment, you might replace `raise` with `return ""`
    # after you've confirmed all errors are handled or understood.


# Example Usage (for testing) - keep this block as is or remove if not using